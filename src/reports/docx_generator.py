"""
DOCX Report Generator (MAINT-012)

Generates professional DOCX reports from analysis-state.json.
"""

import json
import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger("nexus.reports")


class ReportGenerator:
    """Generate professional DOCX reports from analysis-state.json."""

    def generate_rebuild_report(self, analysis_state_path: str, output_path: str) -> str:
        """Generate rebuild analysis report.

        Args:
            analysis_state_path: Path to analysis-state.json
            output_path: Path where to save the DOCX report

        Returns:
            Path to generated report
        """
        # Load analysis state
        with open(analysis_state_path) as f:
            state = json.load(f)

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading(state["projectName"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Executive Summary
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(f"Analysis Date: {state['analyzedAt']}")
        doc.add_paragraph(f"Total Findings: {state['summary']['totalFindings']}")

        # Summary table
        table = doc.add_table(rows=5, cols=2)
        table.style = "Light Grid Accent 1"

        # Header row
        table.rows[0].cells[0].text = "Severity"
        table.rows[0].cells[1].text = "Count"

        severity_rows = [
            ("CRITICAL", state["summary"]["bySeverity"].get("CRITICAL", 0)),
            ("HIGH", state["summary"]["bySeverity"].get("HIGH", 0)),
            ("MEDIUM", state["summary"]["bySeverity"].get("MEDIUM", 0)),
            ("LOW", state["summary"]["bySeverity"].get("LOW", 0)),
        ]

        for idx, (severity, count) in enumerate(severity_rows, start=1):
            table.rows[idx].cells[0].text = severity
            table.rows[idx].cells[1].text = str(count)

        # Findings by category
        doc.add_page_break()
        doc.add_heading("Findings by Category", 1)

        for category, count in sorted(state["summary"]["byCategory"].items()):
            doc.add_heading(f"{category} ({count} findings)", 2)

            # Filter findings for this category
            category_findings = [
                f for f in state["findings"] if f["category"] == category
            ]

            for finding in category_findings:
                doc.add_heading(
                    f"{finding['id']}: {finding['title']}", 3
                )
                doc.add_paragraph(f"Severity: {finding['severity']}")
                doc.add_paragraph(f"Location: {finding['location']}")
                doc.add_paragraph(finding["description"])
                doc.add_paragraph(f"Remediation:\n{finding['remediation']}")
                doc.add_paragraph(
                    f"Effort: {finding['effort']} ({finding['effort_hours']})"
                )

        # Save
        doc.save(output_path)
        logger.info("Report generated: %s", output_path)
        return output_path
